from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service

def create_webdriver():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--disable-software-rasterizer")
    chrome_options.add_argument("--window-size=1920x1080")

    chrome_path = "/usr/bin/google-chrome"
    driver_path = "/usr/bin/chromedriver"

    chrome_options.binary_location = chrome_path

    service = Service(executable_path=driver_path)
    driver = webdriver.Chrome(service=service, options=chrome_options)

    return driver


# TCS, Tech Mahindra, Mphasis, Zensar, Infosys, Wipro, Persistent, Cognizant
#pip install requests pdfplumber streamlit selenium reportlab

import base64
import os
import re
import time
import requests
import pdfplumber
import streamlit as st
from io import BytesIO
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By

# --- Streamlit Setup ---
st.set_page_config(page_title="BFSI Highlights Extractor", layout="centered")
logo_path = os.path.join("assets", "hcl.png")

# Add background and box styling
st.markdown(
    f"""
    <style>
    .top-right-logo {{
        position: fixed;
        top: 70px;
        right: 20px;
        width: 140px;
        z-index: 100;
    }}
    </style>
    <img src="data:image/png;base64,{base64.b64encode(open(logo_path, "rb").read()).decode()}" class="top-right-logo">
    """,
    unsafe_allow_html=True
)
st.markdown(
    """
    <style>
    .white-box {
        background-color: rgba(255, 255, 255, 0.95);
        padding: 2rem;
        border-radius: 12px;
        margin-top: 1.5rem;
        box-shadow: 0 0 15px rgba(0, 0, 0, 0.1);
    }
    </style>
    """,
    unsafe_allow_html=True
)
background_url = "https://images.rawpixel.com/image_800/cHJpdmF0ZS9sci9pbWFnZXMvd2Vic2l0ZS8yMDI0LTAyL3Jhd3BpeGVsX29mZmljZV8zNF9jbG9zZXVwX3Bob3RvX2JsYW5rX3Bvc3Rlcl9tb2NrdXB3aGl0ZV93YWxsY181ODljN2QxNi1kM2YwLTQyYTItOTQ4ZS1hYzc2M2UyMjNhNmRfMS5qcGc.jpg"
st.markdown(
    f"""
    <style>
    .stApp {{
        background: linear-gradient(rgba(255,255,255,0.85), rgba(255,255,255,0.85)),url("https://images.rawpixel.com/image_800/cHJpdmF0ZS9sci9pbWFnZXMvd2Vic2l0ZS8yMDI0LTAyL3Jhd3BpeGVsX29mZmljZV8zNF9jbG9zZXVwX3Bob3RvX2JsYW5rX3Bvc3Rlcl9tb2NrdXB3aGl0ZV93YWxsY181ODljN2QxNi1kM2YwLTQyYTItOTQ4ZS1hYzc2M2UyMjNhNmRfMS5qcGc.jpg");
        background-image: url("{background_url}");
        background-size: cover;
        background-attachment: fixed;
        background-repeat: no-repeat;
        background-position: center;
    }}
    </style>
    """,
    unsafe_allow_html=True
)
st.title("BFSI Highlights Extractor")
# --- Company Selection with Custom Dropdown-Like Logo Items ---
def get_base64_image(path):
    with open(path, "rb") as image_file:
        encoded = base64.b64encode(image_file.read()).decode()
        ext = os.path.splitext(path)[1][1:]  # e.g. 'png'
        return f"data:image/{ext};base64,{encoded}"

company_logos = {
    "TCS": os.path.join("assets", "tcs.png"),
    "Tech Mahindra": os.path.join("assets", "techm.png"),
    "Mphasis": os.path.join("assets", "mphasis.png"),
    "Infosys": os.path.join("assets", "info.png"),
    "Zensar": os.path.join("assets", "zensar.png"),
    "Wipro": os.path.join("assets", "wipro.png"),
    "Persistent": os.path.join("assets", "persis.png"),
    "Cognizant": os.path.join("assets", "cog.png")
}

# Create display names with logo using markdown HTML
options = list(company_logos.keys())
option_images = [f"<img src='{get_base64_image(company_logos[c])}' style='height:18px; vertical-align:middle;'> <span style='vertical-align:middle;'>{c}</span>" for c in options]
company = st.selectbox("Select a company", options, index=0)

# Show selected company logo
logo_data = get_base64_image(company_logos[company])
st.markdown(f"""
<div style='display: flex; align-items: center;'>
    <h2>Selected Company :</h2>
    <img src='{logo_data}' style='height:100px; margin-right:8px;'>
</div>
""", unsafe_allow_html=True)

# Save in session state
st.session_state.selected_company = company

st.markdown("""
    <style>
        .stRadio > div {
            margin-top: -20px;  /* reduce gap under title */
        }
    </style>
""", unsafe_allow_html=True)

# --- Inputs ---
if company == "Infosys":
    fiscal_year = st.text_input("Enter Financial Year (e.g., 2025-2026)")
    quarter_display_map = {
        "Q1 (Apr-Jun)": "q1",
        "Q2 (Jul-Sep)": "q2",
        "Q3 (Oct-Dec)": "q3",
        "Q4 (Jan-Mar)": "q4"
    }
    quarter_label = st.selectbox("Select Quarter:", list(quarter_display_map.keys()))
    quarter_code = quarter_display_map[quarter_label]

elif company == "Zensar":
    fiscal_year = st.text_input("Enter Fiscal Year (e.g., 26 for FY26)")
    quarter_code = st.selectbox("Select Quarter", ["Q1", "Q2", "Q3", "Q4"])
    quarter_label = quarter_code  # Display as-is

elif company == "Wipro":
    fiscal_year = st.text_input("Enter Financial Year (e.g., 2025-2026)")
    quarter_code = st.selectbox("Select Quarter", ["Q1", "Q2", "Q3", "Q4"])
    quarter_label = quarter_code

elif company == "Persistent":
    fiscal_year = st.text_input("Enter Fiscal Year (e.g., 2026 for FY26):")
    quarter_code = st.selectbox("Select Quarter", ["Q1", "Q2", "Q3", "Q4"])
    quarter_label = quarter_code

elif company == "Cognizant":
    fiscal_year = st.text_input("Enter Fiscal Year (e.g., 2025 for FY25):")
    quarter_code = st.selectbox("Select Quarter", ["Q1", "Q2", "Q3", "Q4"])
    quarter_label = quarter_code

else:
    fiscal_year = st.text_input("Enter Financial Year (e.g., 2025-26)")
    quarter_display_map = {
        "Q1 (Apr-Jun)": "q1",
        "Q2 (Jul-Sep)": "q2",
        "Q3 (Oct-Dec)": "q3",
        "Q4 (Jan-Mar)": "q4"
    }
    quarter_code_map_tcs = {
        "Q1 (Apr-Jun)": "quarter1",
        "Q2 (Jul-Sep)": "quarter2",
        "Q3 (Oct-Dec)": "quarter3",
        "Q4 (Jan-Mar)": "quarter4"
    }
    quarter_label = st.selectbox("Select Quarter:", list(quarter_display_map.keys()))
    quarter_code = quarter_display_map[quarter_label]
    quarter_input_tcs = quarter_code_map_tcs[quarter_label]

# --- Keywords ---
keywords = [
    "bfsi", "bank", "banking", "insurance", "financial services", "securities", "retail",
    "investment", "finance", "financial", "asset management", "wealth management", "capital markets",
    "payments", "cards and payments", "brokerage", "trading", "fintech", "treasury", "insurer",
    "reinsurance", "reinsurer", "insurtech", "mortgage", "lender", "lending",
    "custody and fund administration", "custodian","debt"
]
keywords_lower = [k.lower() for k in keywords]

def matches_keywords(line):
    return any(keyword in line.lower() for keyword in keywords)

def highlight_keywords(text):
    for kw in keywords:
        pattern = re.compile(rf"({re.escape(kw)})", re.IGNORECASE)
        text = pattern.sub(r"<b>\1</b>", text)
    return text


def create_docx(highlights, fy, quarter, company):
    doc = Document()
    doc.add_heading(f"{company} BFSI Highlights – FY {fy}, {quarter}", level=1)
    for idx, sentence in enumerate(highlights, 1):
        p = doc.add_paragraph(f"{idx}. ", style="List Number")
        words = re.split(r'(\W+)', sentence)
        for word in words:
            run = p.add_run(word)
            if word.lower() in keywords_lower:
                run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Extractors ---
def extract_tcs(fy_input, quarter_input_tcs):
    highlights = []
    url = f"https://www.tcs.com/investor-relations/financial-statements#year={fy_input}&quarter={quarter_input_tcs}"

    options = Options()
    options.add_argument('--headless')
    options.add_argument('--disable-gpu')
    options.add_argument('--no-sandbox')
    driver = webdriver.Chrome(options=options)
    driver.get(url)
    time.sleep(7)

    pdf_url = None
    links = driver.find_elements(By.TAG_NAME, "a")
    for link in links:
        text = link.text.strip().lower()
        href = link.get_attribute("href")
        if "press release" in text and "usd" in text and href and href.endswith(".pdf"):
            pdf_url = href
            break
    driver.quit()

    if not pdf_url:
        return None, url

    response = requests.get(pdf_url)
    with open("tcs.pdf", "wb") as f:
        f.write(response.content)

    with pdfplumber.open("tcs.pdf") as pdf:
        full_text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                full_text += " ".join([line.strip() for line in page_text.split("\n") if not re.match(r'^\s*Page\s+\d+\s+of\s+\d+\s*$', line)])

    match = re.search(r'(Key Highlights|Key Wins)(.*?)(Customer Speak)', full_text, re.DOTALL | re.IGNORECASE)
    if not match:
        return [], pdf_url

    highlight_block = match.group(2)
    collapsed_block = re.sub(r'\s{2,}', ' ', highlight_block.strip())
    collapsed_block = re.sub(r'\(.*?ranks.*?\)', '', collapsed_block, flags=re.IGNORECASE)
    sentences = re.split(r'(?<=[.])\s+', collapsed_block)
    for sentence in sentences:
        if matches_keywords(sentence):
            highlights.append(sentence.strip())

    return highlights, pdf_url

def extract_techm(fy_input, quarter_code):
    fy_match = re.match(r'^20(\d{2})-?(\d{2})$', fy_input)
    if not fy_match:
        return None, ""
    fy_short = fy_match.group(2)
    pdf_url = f"https://insights.techmahindra.com/investors/tml-{quarter_code}-fy-{fy_short}-press-release.pdf"
    try:
        response = requests.get(pdf_url)
        if response.status_code != 200:
            return None, pdf_url
        with open("techm_press_release.pdf", "wb") as f:
            f.write(response.content)

        highlights = []
        with pdfplumber.open("techm_press_release.pdf") as pdf:
            full_text = " ".join(
                " ".join(re.sub(r'^\s*•\s*', '', line).strip() for line in page.extract_text().split("\n"))
                for page in pdf.pages if page.extract_text()
            )

        match = re.search(r'(Key Deal Wins|Key Wins)(.*?)(Business Highlights)', full_text, re.DOTALL | re.IGNORECASE)
        if not match:
            return [], pdf_url

        win_block = match.group(2)
        cleaned_block = re.sub(r'\s{2,}', ' ', win_block.strip())
        sentences = re.split(r'(?<=[.])\s+', cleaned_block)
        return [s.strip() for s in sentences if matches_keywords(s)], pdf_url
    except Exception as e:
        return None, str(e)

def extract_mphasis(fy_input, quarter_code):
    end_year = "20" + fy_input.split("-")[1]
    url = f"https://www.mphasis.com/content/dam/mphasis-com/global/en/investors/financial-results/{end_year}/{quarter_code}-earnings-press-release.pdf"
    response = requests.get(url)
    if response.status_code != 200:
        return None, url
    with open("mphasis.pdf", "wb") as f:
        f.write(response.content)
    with pdfplumber.open("mphasis.pdf") as pdf:
        full_text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    match = re.search(r"(deal wins\s*:?.*?)(awards and recognitions|recognitions and analyst positioning)", full_text, re.IGNORECASE | re.DOTALL)
    if not match:
        return [], url
    section = re.sub(r'[\u2022\u2023\u25AA\u25CF\u2013\u2014\-]', '\n', match.group(1))
    fragments = re.split(r'(?<=[.!?])\s+(?=[A-Z])', re.sub(r'\s+', ' ', section))
    return [line.strip() for line in fragments if matches_keywords(line)], url

def extract_infosys(fiscal_year, quarter):
    base_url = "https://www.infosys.com/investors/reports-filings/quarterly-results"
    pdf_url = f"{base_url}/{fiscal_year}/{quarter}/documents/ifrs-usd-press-release.pdf"
    try:
        response = requests.get(pdf_url)
        if response.status_code != 200:
            return None, pdf_url

        with pdfplumber.open(BytesIO(response.content)) as pdf:
            full_text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
        full_text = re.sub(r"Page \d+ of \d+", "", full_text).replace("•", "")
        full_text = re.sub(r"\s+", " ", full_text)

        match = re.search(r"Client wins\s*&\s*Testimonials(.*?)(Recognitions\s*&\s*Awards|Recognitions|Awards)", full_text, re.IGNORECASE | re.DOTALL)
        if not match:
            return [], pdf_url

        section = match.group(1).strip()
        sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z])", section)
        return [highlight_keywords(s.strip()) for s in sentences if matches_keywords(s)], pdf_url
    except Exception as e:
        return None, str(e)

def extract_zensar(fy_input, quarter_code):
    base = "https://www.zensar.com/sites/default/files/investor/analyst-meet/"
    suffixes = [
        f"Zensar-{quarter_code}FY{fy_input}-Press-release.pdf",
        f"Zensar-{quarter_code}FY{fy_input}-Press-Release.pdf",
        f"Zensar-{quarter_code}FY{fy_input}Press-Release.pdf",
        f"Zensar-{quarter_code}FY{fy_input}Press-release.pdf",
    ]
    pdf_url = None
    for suffix in suffixes:
        url = base + suffix
        if requests.head(url).status_code == 200:
            pdf_url = url
            break
    if not pdf_url:
        return None, None

    response = requests.get(pdf_url)
    if response.status_code != 200:
        return None, pdf_url

    def combine_multiline_points(lines):
        points, current = [], ""
        for line in lines:
            if re.match(r"^[•\-\d\.]{0,3}\s*[A-Z]", line):
                if current:
                    points.append(current.strip())
                current = line
            else:
                current += " " + line
        if current:
            points.append(current.strip())
        return points

    with pdfplumber.open(BytesIO(response.content)) as pdf:
        full_text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

    start_match = re.search(r"Significant Wins", full_text, re.IGNORECASE)
    end_matches = [re.search(r"Awards and Recognitions", full_text, re.IGNORECASE),
                   re.search(r"Corporate Excellence Snapshot", full_text, re.IGNORECASE)]
    end_pos = min([m.start() for m in end_matches if m], default=None)

    if not start_match or not end_pos:
        return [], pdf_url

    segment = full_text[start_match.end():end_pos]
    lines = [line.strip() for line in segment.split("\n") if line.strip()]
    points = combine_multiline_points(lines)

    matches = []
    for point in points:
        cleaned = re.sub(r"^[•\-\d\.\s]+", "", point)
        if matches_keywords(cleaned):
            matches.append(cleaned.strip())

    return matches, pdf_url

import fitz  # PyMuPDF for Wipro

def extract_wipro(fiscal_year, quarter):
    quarters = {"Q1": "q1", "Q2": "q2", "Q3": "q3", "Q4": "q4"}
    q_lower = quarters[quarter]
    fy = fiscal_year.replace("-", "")
    pdf_url = f"https://www.wipro.com/content/dam/nexus/en/investor/quarterly-results/{fiscal_year.lower()}/{q_lower}fy{fy[-2:]}/press-release-{q_lower}fy{fy[-2:]}.pdf"

    try:
        response = requests.get(pdf_url)
        if response.status_code != 200:
            return None, pdf_url
        doc = fitz.open(stream=response.content, filetype="pdf")
        full_text = "\n".join(page.get_text() for page in doc)
        doc.close()

        # Extract relevant section
        start_patterns = [
            r"(?i)Highlights of Strategic Deal Wins",
            r"(?i)IT Services\s*[-–]\s*Large deals"
        ]
        end_patterns = [
            r"(?i)Analyst Recognition",
            r"(?i)About Key Metrics and Non-GAAP Financial Measures"
        ]

        start_match = next((re.search(p, full_text) for p in start_patterns if re.search(p, full_text)), None)
        if not start_match:
            return [], pdf_url
        start_idx = start_match.end()

        remaining_text = full_text[start_idx:]
        end_match = next((re.search(p, remaining_text) for p in end_patterns if re.search(p, remaining_text)), None)
        section = remaining_text[:end_match.start()] if end_match else remaining_text

        # Break into sentences and clean
        raw_sentences = re.split(r'(?<=[.?!])\s+(?=[A-Z])', section)
        result_sentences = []
        for sent in raw_sentences:
            cleaned = sent.strip()
            if re.match(r'^\s*\d{1,2}[\.\)]\s*$', cleaned):
                continue
            if not cleaned:
                continue
            cleaned = re.sub(r'^\s*\d{1,2}[\.\)]\s*', '', cleaned)
            cleaned = re.sub(r'^[•\-]+\s*', '', cleaned)

            if matches_keywords(cleaned):
                result_sentences.append(cleaned)

        return result_sentences, pdf_url

    except Exception as e:
        return None, str(e)

def extract_persistent(fy_input, quarter_code):
    quarter_month_map = {"Q1": "07", "Q2": "10", "Q3": "01", "Q4": "04"}
    month = quarter_month_map[quarter_code]
    q_lower = quarter_code.lower()
    q_upper = quarter_code.upper()
    fy_suffix = f"fy{str(fy_input)[-2:]}"
    year_prefix = str(int(fy_input) - 1) if quarter_code in ["Q1", "Q2"] else str(fy_input)

    urls = [
        f"https://www.persistent.com/wp-content/uploads/{year_prefix}/{month}/press-release-{q_lower}{fy_suffix}.pdf",
        f"https://www.persistent.com/wp-content/uploads/{year_prefix}/{month}/Press-Release-{q_upper}{fy_suffix.upper()}.pdf",
        f"https://www.persistent.com/wp-content/uploads/{year_prefix}/{month}/press-release-{q_upper}{fy_suffix}.pdf",
        f"https://www.persistent.com/wp-content/uploads/{year_prefix}/{month}/Press-Release-{q_lower}{fy_suffix}.pdf"
    ]

    for url in urls:
        try:
            resp = requests.get(url)
            if resp.status_code == 200:
                with pdfplumber.open(BytesIO(resp.content)) as pdf:
                    full_text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
                match = re.search(r'Banking, Financial Services & Insurance(.*?)Healthcare & Life Sciences', full_text, re.IGNORECASE | re.DOTALL)
                if not match:
                    return ["⚠️ BFSI section not found between 'Banking, Financial Services & Insurance' and 'Healthcare & Life Sciences'."], url

                section = match.group(1).replace("\\", ". ")
                sentences = re.split(r'(?<=[.!?])\s+', section.strip())
                matches = [s.strip() for s in sentences if matches_keywords(s)]
                return matches or ["⚠️ No matching BFSI-related sentences found."], url
        except Exception as e:
            return None, f"Error: {e}"

    return None, f"No valid PDF found for FY{fy_input}, {quarter_code}."

def extract_cognizant(fy_input, quarter_code):
    BFSI_KEYWORDS = [
        "bfsi", "bank", "banking", "insurance", "financial services", "securities", "retail",
        "investment", "finance", "financial", "asset management", "wealth management", "capital markets",
        "payments", "cards and payments", "brokerage", "trading", "fintech", "treasury", "insurer",
        "reinsurance", "reinsurer", "insurtech", "mortgage", "lender", "lending",
        "custody and fund administration", "custodian", "debt"
    ]

    quarter_map = {
        "Q1": "q1", "Q2": "q2", "Q3": "q3", "Q4": "q4"
    }
    qstr = quarter_map[quarter_code]
    fy_short = str(int(fy_input) % 100).zfill(2)

    url = f"https://cognizant.q4cdn.com/123993165/files/doc_earnings/{fy_input}/{qstr}/earnings-result/{quarter_code}-{fy_short}-Earnings-Press-Release.pdf"
    response = requests.get(url)
    if response.status_code != 200:
        return None, url

    with pdfplumber.open(BytesIO(response.content)) as pdf:
        text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    text = re.sub(r'\s+', ' ', text)

    start_patterns = [
        r"(Client Announcements)",
        r"(Client Wins)",
        r"(Select Client and Partnership Announcements)"
    ]
    end_patterns = [
        r"(Platform Enhancements and Partnerships)",
        r"(Select Analyst Ratings, Company Recognition and Announcements)",
        r"(Select Analyst Ratings and Company Recognition)"
    ]

    start_idx = -1
    for pat in start_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            start_idx = m.end()
            break
    if start_idx == -1:
        return [], url

    end_idx = len(text)
    for pat in end_patterns:
        m = re.search(pat, text[start_idx:], re.IGNORECASE)
        if m:
            end_idx = start_idx + m.start()
            break

    section = text[start_idx:end_idx]
    chunks = re.split(r'•|\.\s+(?=[A-Z])', section)

    matches = []
    for chunk in chunks:
        chunk = chunk.strip()
        if any(re.search(rf'\b{re.escape(k)}\b', chunk, re.IGNORECASE) for k in BFSI_KEYWORDS):
            matches.append(chunk)
    return matches, url


# --- Run Button ---
if st.button("Get Highlights"):
    if not fiscal_year or \
   (company in ["Infosys", "Wipro"] and not re.match(r'^\d{4}-\d{4}$', fiscal_year)) or \
   (company in ["TCS", "Tech Mahindra", "Mphasis"] and not re.match(r'^\d{4}-\d{2}$', fiscal_year)) or \
   (company == "Zensar" and not re.match(r'^\d{2}$', fiscal_year)):
        st.error("❌ Invalid fiscal year format.")
        st.stop()

    highlights, source_url = [], ""
    if company == "TCS":
        highlights, source_url = extract_tcs(fiscal_year, quarter_input_tcs)
    elif company == "Tech Mahindra":
        highlights, source_url = extract_techm(fiscal_year, quarter_code)
    elif company == "Mphasis":
        highlights, source_url = extract_mphasis(fiscal_year, quarter_code)
    elif company == "Infosys":
        highlights, source_url = extract_infosys(fiscal_year, quarter_code)
    elif company == "Zensar":
        highlights, source_url = extract_zensar(fiscal_year, quarter_code)
    elif company == "Wipro":
        highlights, source_url = extract_wipro(fiscal_year, quarter_code)
    elif company == "Persistent":
        if not re.match(r'^\d{4}$', fiscal_year):
            st.error("❌ Invalid fiscal year format for Persistent. Use e.g., 2026.")
            st.stop()
        highlights, source_url = extract_persistent(fiscal_year, quarter_code)
    elif company == "Cognizant":
        if not re.match(r'^\d{4}$', fiscal_year):
            st.error("❌ Invalid fiscal year format for Cognizant. Use e.g., 2025.")
            st.stop()
        highlights, source_url = extract_cognizant(fiscal_year, quarter_code)



    if highlights is None:
        st.error(f"❌ PDF not found or error: {source_url}")
    elif highlights:
        with st.container():
            html = f"""
            <div class="white-box">
                <h3>BFSI Highlights from {company} (FY {fiscal_year}, {quarter_label})</h3>
            """
            for idx, line in enumerate(highlights, 1):
                highlighted = highlight_keywords(line)
                html += f"<p><b>{idx}.</b> {highlighted}</p>"

            html += "</div>"  # Close white-box
            st.markdown(html, unsafe_allow_html=True)

            # Show download button inside Streamlit block
            docx_file = create_docx(highlights, fiscal_year, quarter_label, company)
            st.download_button(
                label="📥 Download Highlights as Word Document",
                data=docx_file,
                file_name=f"{company}_BFSI_Highlights_{fiscal_year}_{quarter_code}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    else:
        st.warning("No BFSI-related highlights found.")

